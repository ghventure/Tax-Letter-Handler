import csv
from docx import Document
from docx.shared import Pt
from datetime import date
import os

#Get today's date
today = date.today()
formattedDate = today.strftime("%B %d, %Y")

def main():

    #Figure out the "host" path where template files and program exe are located;
    #program will return here after creating each letter/envelope in the new folder
    hostPath = os.getcwd()
    print("Host path is: " + hostPath)
    #Create a new folder for all generated letters and envelopes (if it doesn't already exist)
    #newPath = r'C:\Users\jackc\OneDrive\Desktop\code\Generated_Tax_Letters_and_Envelopes'
    #if not os.path.exists(newPath):
    #    os.makedirs(newPath)
    #os.chdir(newPath)

    #Take input from user regarding their intentions with the app
    print ("Hello! Welcome to the Tax Letter and Envelope Generator!")
    print ("This program will read donor information from a CSV file and " \
    "generate personalized tax letters \nand envelopes for each donor.")
    print ("You can also use this program to create a single letter and envelope. \n")

    #Ask the user if they want to use a CSV file or create a single letter/envelope
    useCSV = input("Would you like to use a CSV file to generate multiple letters and envelopes? (yes/no): ").strip().lower()
    if useCSV == 'yes':
        csvFileName = input("Please enter the case-sensitive name of your .csv file (including the .csv extension): ")
        
        

        #Create a new folder for all generated letters and envelopes (if it doesn't already exist)
        hostPathString = str(hostPath)
        newPath = hostPathString + r'\Generated_Tax_Letters_and_Envelopes'
        if not os.path.exists(newPath):
            os.makedirs(newPath)
        #os.chdir(newPath)

        #Call functions to create tax letters and envelopes from the CSV file
        #print("We are currently in " + os.getcwd())
        #print("Changing path to host path: " + hostPath)
        #print("\n\n\n")
        os.chdir(hostPath)
        #print("We are currently in " + os.getcwd())
        with open(csvFileName, mode='r') as file:
            reader = csv.reader(file)
            for row in reader:
                donorFirstName = row[0]
                donorLastName = row[1]
                donorAddress = row[2]
                donorCityState = row[3]
                donorZip = row[4]
                donorAmount = row[5]      
                createTaxLettersFromCSV(hostPath, newPath, donorFirstName, donorLastName, donorAddress, donorCityState, donorZip, donorAmount)
                createEnvelopesFromCSV(hostPath, newPath, donorFirstName, donorLastName, donorAddress, donorCityState, donorZip)

    #Handle case where user will input data manually for a single letter/envelope at a time
    if useCSV == 'no':
        #Need to establish new path for generated letters/envelopes, don't want to make new folder unless we get input
        hostPathString = str(hostPath)
        newPath = hostPathString + r'\Generated_Tax_Letters_and_Envelopes'
        if not os.path.exists(newPath):
            os.makedirs(newPath)

        print("Okay, let's get started on a single letter and envelope.")
        #Collect donor information via input prompts
        donorFirstName = input("Please enter the donor's first name: ")
        donorLastName = input("Please enter the donor's last name: ")
        donorAddress = input("Please enter the donor's address: ")
        donorCityState = input("Please enter the donor's city and state: ")            
        donorZip = input("Please enter the donor's zip code: ")
        donorAmount = input("Please enter the donation amount (Include dollar sign $): ")
        createTaxLettersFromCSV(hostPath, newPath, donorFirstName, donorLastName, donorAddress, donorCityState, donorZip, donorAmount)
        createEnvelopesFromCSV(hostPath, newPath, donorFirstName, donorLastName, donorAddress, donorCityState, donorZip)


#Ceate tax letters for each donor in the CSV file
def createTaxLettersFromCSV(hostPath, newPath, donorFirstName, donorLastName, donorAddress, donorCityState, donorZip, donorAmount):

        #Create a new document for each donor based on the template
        print("We are currently in " + os.getcwd())
        newTaxLetter = Document('_taxTemplateForCSV.docx')

        #set up the font style and size
        style = newTaxLetter.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)
        #Replace placeholders in the document with actual donor information
        for paragraph in newTaxLetter.paragraphs:
            if '{FirstName}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{FirstName}', donorFirstName)
                paragraph.style = newTaxLetter.styles['Normal']
            if '{LastName}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{LastName}', donorLastName)
                paragraph.style = newTaxLetter.styles['Normal']
            if '{Address}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{Address}', donorAddress)
                paragraph.style = newTaxLetter.styles['Normal']
            if '{CityState}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{CityState}', donorCityState)
                paragraph.style = newTaxLetter.styles['Normal']
            if '{Zip}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{Zip}', donorZip)
                paragraph.style = newTaxLetter.styles['Normal']
            if '{Amount}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{Amount}', donorAmount)
                paragraph.style = newTaxLetter.styles['Normal']
            if '{Date}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{Date}', formattedDate)
                paragraph.style = newTaxLetter.styles['Normal']

            #Save the new document with a unique name for each donor
        print("Changing path to new directory: " + newPath)
        os.chdir(newPath)
        #os.makedirs(newPath, exist_ok=True)
        #newTaxLetter.save(newPath, f'tax{donorFirstName}{donorLastName}_{today}.docx')
        newTaxLetter.save(f'tax{donorFirstName}{donorLastName}_{today}.docx')

        os.chdir(hostPath)

#create Envelopes for each donor in the CSV file
def createEnvelopesFromCSV(hostPath, newPath, donorFirstName, donorLastName, donorAddress, donorCityState, donorZip):

    #Move to the directory where the template is located
    #os.chdir(hostPath)
    print("We are currently in " + os.getcwd())

    #Create a new document for each envelope based on the template
    newEnvelope = Document('_envelopeTemplateForCSV.docx')
            
    #Replace placeholders in the document with actual donor information
    for paragraph in newEnvelope.paragraphs:
        if '{FirstName}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{FirstName}', donorFirstName)
        if '{LastName}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{LastName}', donorLastName)
        if '{Address}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{Address}', donorAddress)
        if '{CityState}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{CityState}', donorCityState)
        if '{Zip}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{Zip}', donorZip)
    #Check tables too
    for table in newEnvelope.tables:
        for row in table.rows:
            for cell in row.cells:
                if '{FirstName}' in cell.text:
                    cell.text = cell.text.replace('{FirstName}', donorFirstName)
                if '{LastName}' in cell.text:
                    cell.text = cell.text.replace('{LastName}', donorLastName)
                if '{Address}' in cell.text:
                    cell.text = cell.text.replace('{Address}', donorAddress)
                if '{CityState}' in cell.text:
                    cell.text = cell.text.replace('{CityState}', donorCityState)
                if '{Zip}' in cell.text:
                    cell.text = cell.text.replace('{Zip}', donorZip)
                            
    #Save the new document with a unique name for each donor
    os.chdir(newPath)
    #os.makedirs(newPath, exist_ok=True)
    #newEnvelope.save(newPath, f"envelope{donorFirstName}{donorLastName}.docx")
    newEnvelope.save(f"envelope{donorFirstName}{donorLastName}.docx")

    os.chdir(hostPath)

if __name__ == "__main__":
    main()